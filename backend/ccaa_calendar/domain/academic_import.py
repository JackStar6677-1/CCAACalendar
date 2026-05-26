from __future__ import annotations

import csv
import re
from dataclasses import asdict, dataclass
from datetime import UTC, datetime, time, timedelta
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

SUPPORTED_IMPORT_EXTENSIONS = {".csv", ".txt", ".xlsx", ".docx", ".pdf"}

MONTHS_ES = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "setiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}

DATE_PATTERNS = [
    re.compile(r"\b(?P<day>\d{1,2})[/-](?P<month>\d{1,2})[/-](?P<year>\d{2,4})\b"),
    re.compile(r"\b(?P<year>\d{4})-(?P<month>\d{1,2})-(?P<day>\d{1,2})\b"),
    re.compile(
        r"\b(?P<day>\d{1,2})\s+de\s+"
        r"(?P<month_name>enero|febrero|marzo|abril|mayo|junio|julio|agosto|"
        r"septiembre|setiembre|octubre|noviembre|diciembre)"
        r"(?:\s+de)?\s+(?P<year>\d{4})\b",
        re.IGNORECASE,
    ),
]

CATEGORY_KEYWORDS = {
    "feriado": ("feriado", "irrenunciable"),
    "academico": ("semestre", "catedra", "cátedra", "examen", "evaluacion", "evaluación"),
    "espacio": ("auditorio", "sala", "reserva", "campus"),
    "centro": ("asamblea", "reunion", "reunión", "centro de estudiantes", "ce "),
}


@dataclass
class AcademicImportCandidate:
    title: str
    starts_at: datetime
    ends_at: datetime
    category: str
    description: str
    source_line: str
    confidence: float

    def to_dict(self) -> dict:
        payload = asdict(self)
        payload["starts_at"] = self.starts_at.isoformat()
        payload["ends_at"] = self.ends_at.isoformat()
        return payload


def extension_for_filename(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def extract_text_lines(filename: str, content: bytes) -> list[str]:
    """Convierte archivos comunes de calendario en lineas planas auditables."""
    extension = extension_for_filename(filename)
    if extension not in SUPPORTED_IMPORT_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_IMPORT_EXTENSIONS))
        raise ValueError(f"Formato no soportado. Usa uno de estos: {supported}.")

    if extension == ".xlsx":
        return _lines_from_xlsx(content)
    if extension == ".docx":
        return _lines_from_docx(content)
    if extension == ".pdf":
        return _lines_from_pdf(content)
    if extension == ".csv":
        return _lines_from_csv(content)
    return _lines_from_text(content)


def parse_academic_calendar(filename: str, content: bytes) -> dict:
    """Extrae candidatos de eventos sin escribir en la base de datos."""
    lines = extract_text_lines(filename, content)
    candidates: list[AcademicImportCandidate] = []

    for line in lines:
        candidate = _candidate_from_line(line)
        if candidate:
            candidates.append(candidate)

    # Los calendarios institucionales suelen repetir hitos entre portada, tabla y anexos.
    # Deduplizamos por titulo+fecha para evitar que una aprobacion cree clones.
    seen: set[tuple[str, str]] = set()
    unique_candidates = []
    for candidate in candidates:
        key = (candidate.title.casefold(), candidate.starts_at.date().isoformat())
        if key in seen:
            continue
        seen.add(key)
        unique_candidates.append(candidate)

    return {
        "filename": filename,
        "line_count": len(lines),
        "candidates": [candidate.to_dict() for candidate in unique_candidates],
        "warnings": _import_warnings(lines, unique_candidates),
    }


def _lines_from_text(content: bytes) -> list[str]:
    text = _decode_text(content)
    return _clean_lines(text.splitlines())


def _lines_from_csv(content: bytes) -> list[str]:
    text = _decode_text(content)
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    reader = csv.reader(StringIO(text), dialect)
    return _clean_lines(" | ".join(cell.strip() for cell in row if cell.strip()) for row in reader)


def _lines_from_xlsx(content: bytes) -> list[str]:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell not in (None, "")]
            if cells:
                lines.append(" | ".join(cells))
    return _clean_lines(lines)


def _lines_from_docx(content: bytes) -> list[str]:
    document = Document(BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return _clean_lines(lines)


def _lines_from_pdf(content: bytes) -> list[str]:
    reader = PdfReader(BytesIO(content))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        lines.extend(text.splitlines())
    return _clean_lines(lines)


def _clean_lines(lines) -> list[str]:
    cleaned = []
    for line in lines:
        normalized = " ".join(str(line).replace("\xa0", " ").split())
        if len(normalized) >= 6:
            cleaned.append(normalized)
    return cleaned


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _candidate_from_line(line: str) -> AcademicImportCandidate | None:
    """Transforma una linea con fecha reconocible en un evento candidato."""
    parsed_date = _find_date(line)
    if not parsed_date:
        return None

    date_value, date_text = parsed_date
    title = _title_from_line(line, date_text)
    if len(title) < 3:
        title = "Hito académico importado"

    category = _category_for_line(line)
    # Si el documento no trae hora, usamos 09:00 como hora neutra de trabajo.
    starts_at = datetime.combine(date_value, time(9, 0), tzinfo=UTC)
    ends_at = starts_at + timedelta(hours=1)

    return AcademicImportCandidate(
        title=title[:220],
        starts_at=starts_at,
        ends_at=ends_at,
        category=category,
        description=f"Importado desde calendario académico: {line[:500]}",
        source_line=line[:800],
        confidence=_confidence_for_line(line, category),
    )


def _find_date(line: str) -> tuple[datetime.date, str] | None:
    for pattern in DATE_PATTERNS:
        match = pattern.search(line)
        if not match:
            continue
        groups = match.groupdict()
        year = int(groups["year"])
        if year < 100:
            year += 2000
        month = (
            int(groups["month"])
            if groups.get("month")
            else MONTHS_ES[groups["month_name"].lower()]
        )
        day = int(groups["day"])
        try:
            return datetime(year, month, day, tzinfo=UTC).date(), match.group(0)
        except ValueError:
            return None
    return None


def _title_from_line(line: str, date_text: str) -> str:
    without_date = line.replace(date_text, " ")
    without_separators = re.sub(r"\s*[|:;-]\s*", " ", without_date)
    return " ".join(without_separators.split()).strip(" .,-")


def _category_for_line(line: str) -> str:
    normalized = line.casefold()
    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(keyword in normalized for keyword in keywords):
            return category
    return "academico"


def _confidence_for_line(line: str, category: str) -> float:
    score = 0.65
    if category != "academico":
        score += 0.1
    if any(separator in line for separator in ("|", ":", "-")):
        score += 0.1
    if len(line) > 18:
        score += 0.1
    return min(score, 0.95)


def _import_warnings(lines: list[str], candidates: list[AcademicImportCandidate]) -> list[str]:
    warnings = []
    if not lines:
        warnings.append("No se pudo leer texto util del archivo.")
    if not candidates:
        warnings.append("No encontramos fechas reconocibles. Revisa formato o usa CSV/Excel.")
    elif len(candidates) < 3:
        warnings.append("Encontramos pocos hitos; revisa la previsualizacion antes de aprobar.")
    return warnings
